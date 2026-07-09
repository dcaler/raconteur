"""The redline revise path: answer each comment with a minimal, in-place tracked change.

This is what `raconteur paper` does by default when it finds an annotated .docx. The clean
rewrite survives as `paper --resynth`.

Three ideas carry the whole module:

  1. THE ANNOTATION BLOB IS DEAD. `revise.build_revision_context` concatenates every comment
     and tracked change into one string handed to every section with "apply only those
     relevant; ignore the rest". That leaves routing to the model, which is why a comment on
     the Discussion used to rewrite the Methods. Here each comment is anchored to the exact
     sentences it spans, in the paragraph it belongs to, and no other sentence in the
     document is ever shown to the reviser.

  2. THE REVISER RETURNS ONLY THE SENTENCES IT CHANGED, keyed by index. A sentence it does
     not return is copied byte-for-byte — the untouched sentences are literally the original
     strings, so minimality is true by construction rather than hoped for from a diff. That
     also makes the touched set exact, so `minimal_edit_violation` can prove a comment on
     sentence 2 did not rewrite sentence 4.

  3. FAIL CLOSED. Malformed JSON, a dropped citekey, a dropped or invented equation, an
     out-of-scope sentence, an exhausted retry budget — any of these and we write NO tracked
     change and say so in the reply. A broken edit under a reply claiming "done" is the worst
     outcome, worse than a visibly skipped comment.

Guards in Python, judgement in the LLM: everything mechanical is decided here, precisely, and
stated as an imperative. The audit call is left with the one question code cannot answer —
does this edit mean what the comment asked for?
"""

from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

from . import guards, redline
from .brain import Brain
from .config import ProjectConfig
from .log import log

# ── the per-paragraph reviser ────────────────────────────────────────────────

_PARA_REVISE_SYS = """\
You are revising ONE paragraph of a scholarly paper to satisfy a reviewer's comment(s) on it.
The paragraph is given to you as NUMBERED SENTENCES. You return only the sentences you
changed, keyed by their number — never the whole paragraph. A sentence you do not return
survives word for word, which is the point: it keeps its citations, its grounding, and its
evidence intact.

Make the SMALLEST change that fully and genuinely addresses every comment. Revise the
sentence(s) the comment bears on. Leave the rest alone.

CITATIONS — cite ONLY sources in the bibliography below, and ALWAYS as a [@citekey] tag using
the exact key shown: write "[@smith2021]", NEVER "Smith (2021)" or "(Smith, 2021)". An
author-year citation is invisible to the bibliography and silently unverifies the claim.
Every [@citekey] in a sentence you rewrite must survive in your version unless a comment asks
you to remove that source.

PLACEHOLDERS — a token like ⟦m:1⟧ stands for an equation in the original. Reproduce it
exactly, in the sentence whose claim it supports. Never retype an equation as prose, never
move one to another sentence, and never invent a placeholder of your own.

OUTPUT — a single JSON object mapping sentence number to its replacement text. Use null to
delete a sentence. Return nothing else: no prose, no commentary, no code fence.
  {"2": "The revised second sentence [@smith2021].", "5": null}
If no sentence needs to change, return {}."""

_PARA_REVISE_PROMPT = """\
Paper: {title}
Section: {heading}

PARAGRAPH, as numbered sentences. ▶ marks the sentence(s) the reviewer's comment is anchored
to — those are the ones to revise:
{sentences}

REVIEWER COMMENT(S) on this paragraph (address every one):
{comments}
{context_section}{bib_section}
Return the JSON object of changed sentences only."""

_PARA_AUDIT_SYS = """\
You audit ONE revised paragraph of a scholarly paper against the reviewer comment(s) it was
meant to satisfy. Mechanical checks — citation format, dropped citations, equations, which
sentences were touched — have already been made in code and passed; do not repeat them. Judge
only what code cannot: MEANING.

Respond with EXACTLY one of three things, nothing else:
- "OK" — the revision fully and genuinely addresses every comment.
- A line "ROUTE: <class>: <brief reason>" — a comment that CANNOT be satisfied by editing
  this paragraph's prose. <class> is exactly one of:
      section   — asks for a new section or subsection, or material not belonging here
      sources   — asks for literature or citations not present in the bibliography
      evidence  — asks for a result, statistic, or method that does not exist yet
      figure    — asks for a table, chart, or figure
  Do not accept a prose gesture as satisfying such a request.
- Otherwise a numbered list of specific problems: a comment not really addressed, or
  addressed in name only. Quote the text you mean."""

_PARA_AUDIT_PROMPT = """\
Paper: {title}
Section: {heading}

REVIEWER COMMENT(S) the revision must satisfy:
{comments}

ORIGINAL PARAGRAPH:
{paragraph}

REVISED PARAGRAPH (under audit):
{revised}

Judge only against the comment(s): is each fully and genuinely addressed, and is the comment
even satisfiable by editing this paragraph's prose at all? Respond "OK", or
"ROUTE: <class>: <reason>", or a numbered list."""

# What the reviewer is told when a comment cannot be a redline. Answering a request for a
# figure with "this needs sources not in the bibliography" is a false diagnosis: gathering
# papers will never satisfy it. The class is what makes the reply honest.
_ROUTE_CLASSES = ("section", "sources", "evidence", "figure")

_ROUTE_ADVICE = {
    "section": "cannot be a tracked change — it asks for new structure. "
               "Run 'raconteur outline' to revise the structure, then 'raconteur paper'.",
    "sources": "cannot be a tracked change — it asks for literature not in refs.bib. "
               "Run rabbitHole to gather the sources first.",
    "evidence": "cannot be a tracked change — it asks for a result or method that does not "
                "exist. raconteur cannot manufacture evidence: run rayleigh (results) or "
                "raster (methods) first.",
    "figure": "cannot be a tracked change — it asks for a table or figure. "
              "Produce it in rayleigh, then re-render.",
}


def _route_class(verdict: str) -> str:
    """Extract the class from a "ROUTE: <class>: <reason>" verdict."""
    rest = verdict.split(":", 1)[1] if ":" in verdict else ""
    head = rest.strip().split(":", 1)[0].strip().lower()
    return head if head in _ROUTE_CLASSES else "sources"


def _is_ok(verdict: str) -> bool:
    v = verdict.strip().upper()
    return v == "OK" or v.startswith("OK ") or v.startswith("OK.")


# ── sentence-indexed edits ───────────────────────────────────────────────────

def _number_sentences(units: list[str], anchored: set[int]) -> str:
    """Render the paragraph as numbered sentences, marking those a comment bears on."""
    return "\n".join(
        f"{'▶' if i in anchored else ' '} {i + 1}. {u.strip()}"
        for i, u in enumerate(units))


def _apply_sentence_edits(units: list[str], edits: dict) -> str:
    """Rebuild the paragraph from the original units plus the reviser's replacements.

    Every sentence the reviser did not return is copied byte-for-byte. This is what makes the
    sentence-level redline true by construction rather than hoped for from a diff: the
    untouched sentences are literally the original objects.
    """
    out: list[str] = []
    for i, unit in enumerate(units):
        key = str(i + 1)
        if key not in edits:
            out.append(unit)
            continue
        repl = edits[key]
        if repl is None:
            continue  # deleted
        trailing = unit[len(unit.rstrip()):]  # keep the original inter-sentence spacing
        out.append(repl.strip() + trailing)
    return "".join(out)


def _parse_sentence_edits(raw: str, n_units: int) -> tuple[dict, list[str]]:
    """Parse the reviser's JSON, keeping only in-range integer keys. Returns (edits, errors).

    Strict on purpose: a lenient parser that falls back to an empty dict on malformed output
    would look exactly like a well-formed edit of nothing, and we would write no tracked
    change while replying that the comment was addressed.
    """
    m = re.search(r"\{.*\}", raw, re.DOTALL)
    obj = None
    if m:
        try:
            parsed = json.loads(m.group(0))
            obj = parsed if isinstance(parsed, dict) else None
        except json.JSONDecodeError:
            obj = None
    if obj is None:
        return {}, ['Output was not a JSON object. Return only a JSON object mapping '
                    'sentence number to replacement text, e.g. {"2": "…"}.']
    edits, errors = {}, []
    for k, v in obj.items():
        if not str(k).strip().isdigit() or not (1 <= int(k) <= n_units):
            errors.append(f'"{k}" is not a sentence number between 1 and {n_units}.')
            continue
        if v is not None and not isinstance(v, str):
            errors.append(f'The value for sentence {k} must be text or null.')
            continue
        edits[str(int(k))] = v
    return edits, errors


# ── the per-paragraph adversary ──────────────────────────────────────────────

def _para_guard_findings(
    old_text: str, new_text: str, touched: set[int], anchored: set[int],
    n_units: int, kind: str, known: set[str],
) -> list[guards.Finding]:
    """Everything about a paragraph rewrite that Python can decide precisely.

    Note what is NOT here. Density guards (`uncited_paragraphs`, `sparse_paragraphs`) belong
    to the DRAFT phase: on a redline, collateral change is the defect, and a comment asking
    to tighten a sentence does not license injecting citations into it. Note also the
    section-kind gate: a Methods or Results paragraph is grounded in the writeup, not the
    bibliography, so it is not required to carry a citation at all.

    And note what is left for the audit: whether the edit means what the comment asked for.
    Everything else moved into code, where it is exact.
    """
    findings = (guards.author_year_prose(new_text)
                + guards.dropped_citekeys(old_text, new_text)
                + guards.dropped_sentinels(old_text, new_text)
                + guards.invented_sentinels(old_text, new_text)
                + guards.minimal_edit_violation(touched, anchored, n_units))
    if known:
        findings += guards.unresolved_keys(new_text, known)
    if guards.expects_citations(kind) and not guards.CITE_TAG_RE.search(new_text):
        findings.append(guards.Finding(
            "uncited", "paragraph",
            "The paragraph now cites no source — restore a [@citekey] from the bibliography."))
    return findings


def redline_paragraph(
    brain: Brain, title: str, heading: str, paragraph: str, comments: list[str],
    context_section: str, bib_section: str, anchored: set[int], kind: str,
    known: set[str], rounds: int = 2,
) -> tuple[str | None, str]:
    """Rewrite one commented paragraph and hold it to the adversarial bar.

    Returns (new_text, outcome):
      - ("…text…", "edited")   — passes every deterministic guard and the audit.
      - (None, "route:<class>") — a comment a prose edit cannot satisfy; the caller routes it.
      - (None, "skipped")      — no edit could be produced that keeps the paragraph
                                 verifiable. Fail closed: leave the reviewer's paragraph
                                 alone and say so, rather than emit a tracked change that
                                 quietly lost a source or an equation.
    """
    units = guards.sentence_units(paragraph)
    if not units:
        return None, "skipped"
    comment_block = "\n".join(f"- {c}" for c in comments)
    base_prompt = _PARA_REVISE_PROMPT.format(
        title=title, heading=heading,
        sentences=_number_sentences(units, anchored),
        comments=comment_block,
        context_section=context_section,
        bib_section=bib_section,
    )

    critique: str | None = None
    for _ in range(rounds):
        prompt = base_prompt if critique is None else (
            base_prompt + f"\n\nYour previous attempt had these problems — fix every one, "
            f"changing as little else as possible:\n{critique}\n\nReturn the corrected JSON "
            f"object of changed sentences only.")
        try:
            raw = brain.coordinator(prompt, _PARA_REVISE_SYS, num_ctx=16384).strip()
        except Exception as e:
            log(f"[warn] paragraph revise failed ({e}); leaving as-is.")
            return None, "skipped"

        edits, errors = _parse_sentence_edits(raw, len(units))
        if errors:
            critique = "\n".join(f"- {e}" for e in errors)
            continue
        if not edits:
            # The reviser says nothing needs changing. It cannot both leave the paragraph
            # alone and have addressed the comment; let the caller record it as skipped.
            return None, "skipped"

        new_text = _apply_sentence_edits(units, edits)
        touched = {int(k) - 1 for k in edits}

        # Deterministic guards run first: the expensive audit never sees a paragraph that is
        # already broken. Any failure feeds a focused re-revise.
        findings = _para_guard_findings(
            paragraph, new_text, touched, anchored, len(units), kind, known)
        if findings:
            critique = "\n".join(f"- {f.imperative}" for f in findings)
            continue

        # The only question left for the brain: does this edit mean what the comment asked?
        try:
            verdict = brain.coordinator(
                _PARA_AUDIT_PROMPT.format(
                    title=title, heading=heading, comments=comment_block,
                    paragraph=paragraph, revised=new_text),
                _PARA_AUDIT_SYS, num_ctx=16384).strip()
        except Exception as e:
            # Fail closed. The guards prove the text is verifiable, not that it answers the
            # comment — and replying "revised to address this" when nothing checked that
            # claim is exactly the fabricated reply this adversary exists to prevent.
            log(f"[warn] paragraph audit failed ({e}); leaving the paragraph unchanged.")
            return None, "skipped"
        if verdict.upper().startswith("ROUTE"):
            return None, f"route:{_route_class(verdict)}"
        if _is_ok(verdict):
            return new_text, "edited"
        critique = verdict  # audit found problems → another round

    # Rounds exhausted. Fail closed: a paragraph that still trips any guard would emit a
    # tracked change that silently dropped a source or an equation, over a reply claiming the
    # comment was addressed. Leave the paragraph as the reviewer wrote it.
    return None, "skipped"


# ── orchestration ────────────────────────────────────────────────────────────

def _out_path(paper_dir: Path, short_title: str, user_rev: Path) -> Path:
    """A redline is a MINOR version: it keeps the reviewer's datestamp and extends the chain.

    260709_trust_ra_DCR.docx -> 260709_trust_ra_DCR_ra.docx
    """
    from .naming import parse, minor_name
    parsed = parse(user_rev, short_title)
    chain = parsed[1] if parsed else ["ra"]
    datestamp = parsed[0] if parsed else None
    return paper_dir / minor_name(short_title, chain, "docx", datestamp)


def redline_revise(
    project_dir: Path,
    cfg: ProjectConfig,
    brain: Brain,
    paper_dir: Path,
    user_rev: Path,
    litrev: str,
    code: str,
    results: str,
    bib_section: str,
    known: set[str],
) -> tuple[Path, dict[str, str]]:
    """Edit a COPY of the reviewer's .docx in place, one anchored comment at a time.

    Returns (output_path, dispositions) where dispositions maps comment id -> outcome.
    Silence is not a decision: every comment gets a disposition, and every one that could not
    be answered by a tracked change is reported with the reason it could not.
    """
    from docx import Document
    from .paper import _context_for_section

    out = _out_path(paper_dir, cfg.short_title, user_rev)
    shutil.copy2(user_rev, out)

    anchors = redline.comment_anchors(out)
    cmap = redline.comments_by_id(out)
    headings = redline.heading_comments(out)

    if not anchors and not headings:
        log("[warn] no comment anchors found in the revision — nothing to redline")
        return out, {}

    doc = Document(str(out))
    ids = redline.ids_for(doc)
    body = {rec["index"]: rec for rec in redline.body_paragraphs(doc)}
    dispositions: dict[str, str] = {}
    edited = 0

    for anchor in anchors:
        rec = body.get(anchor["index"])
        if rec is None:
            continue
        comments = [cmap[c]["text"] for c in anchor["ids"] if c in cmap]
        if not comments:
            continue
        heading = anchor["heading"] or "Abstract"
        ctx = _context_for_section(anchor["heading"], litrev, code, results)
        context_section = f"\n{ctx}" if ctx else ""

        log(f"[raconteur] redlining '{heading}' para {anchor['index']} "
            f"({len(comments)} comment(s), anchored to sentence(s) "
            f"{[i + 1 for i in anchor['anchored']] or 'all'})…")

        new_text, outcome = redline_paragraph(
            brain, cfg.title, heading, anchor["text"], comments,
            context_section, bib_section, set(anchor["anchored"]),
            anchor["kind"], known,
        )

        for cid in anchor["ids"]:
            dispositions[cid] = outcome

        if outcome == "edited" and new_text:
            if redline.tracked_replace_sentencewise(rec["para"]._p, new_text, redline.AUTHOR, ids):
                edited += 1
                log(f"[raconteur]   → tracked change written")
            else:
                for cid in anchor["ids"]:
                    dispositions[cid] = "skipped"
                log("[warn]   → no textual change; nothing written")
        elif outcome.startswith("route:"):
            cls = outcome.split(":", 1)[1]
            log(f"[warn]   → {_ROUTE_ADVICE.get(cls, 'cannot be a tracked change.')}")
        else:
            log("[warn]   → skipped: no verifiable edit could be produced; "
                "the paragraph is unchanged")

    for h in headings:
        for cid in h["ids"]:
            dispositions[cid] = "route:section"
        log(f"[warn] comment on heading '{h['heading']}' — {_ROUTE_ADVICE['section']}")

    doc.save(str(out))
    log(f"[raconteur] wrote {out.relative_to(project_dir)} "
        f"({edited} paragraph(s) redlined)")
    _report(dispositions, cmap, known, out)
    return out, dispositions


def _report(dispositions: dict[str, str], cmap: dict[str, dict],
            known: set[str], out: Path) -> None:
    """Every comment gets a reply. Silence is not a decision — a comment neither applied nor
    explicitly declined is a defect in the revise pass itself."""
    if not dispositions:
        return
    log("[raconteur] ── comment dispositions ──")
    counts = {"edited": 0, "routed": 0, "declined": 0}
    for cid, outcome in dispositions.items():
        text = cmap.get(cid, {}).get("text", "?")[:60]
        if outcome == "edited":
            verdict = "applied as a tracked change"
            counts["edited"] += 1
        elif outcome.startswith("route:"):
            verdict = _ROUTE_ADVICE.get(outcome.split(":", 1)[1], "routed")
            counts["routed"] += 1
        else:
            verdict = "DECLINED — no verifiable edit could be produced; paragraph unchanged"
            counts["declined"] += 1
        log(f"[raconteur]   [{cid}] {text!r}: {verdict}")
    log(f"[raconteur] {counts['edited']} applied · {counts['routed']} routed · "
        f"{counts['declined']} declined")

    from docx import Document
    doc = Document(str(out))
    md = redline.accepted_markdown(doc)
    log(f"[raconteur] {guards.metrics(md, known)}")

    cited = set(guards.all_citekeys(md))
    if known:
        unresolved = sorted(cited - known)
        if unresolved:
            log(f"[warn] the redlined text cites {len(unresolved)} key(s) with no refs.bib "
                f"entry: {', '.join('[@' + k + ']' for k in unresolved)}")
    # The References list in the .docx was rendered by pandoc/citeproc at draft time. A
    # redline edits prose only, so a newly cited source has no entry there. Rebuilding a
    # styled bibliography inside OOXML is a separate job; say so rather than pretend.
    ref_keys = _references_keys(doc)
    if ref_keys is not None:
        new_keys = sorted(cited - ref_keys)
        if new_keys:
            log(f"[warn] {len(new_keys)} newly cited source(s) have no entry in the rendered "
                f"References list: {', '.join('[@' + k + ']' for k in new_keys)}")
            log("[warn] the References section is not rebuilt by the redline — re-render "
                "with 'raconteur paper --resynth' if the bibliography must be complete")


def _references_keys(doc) -> set[str] | None:
    """Citekeys already present in the rendered References section.

    Returns None when the document has no References section at all — then there is no
    bibliography to have drifted, and warning about it would be noise.

    Best-effort otherwise: pandoc renders entries as prose, so we look for [@key] text, which
    survives when the draft still carries the tags.
    """
    text: list[str] = []
    found = False
    in_refs = False
    for p in doc.paragraphs:
        if redline.is_heading_style(redline._style_name(p)):
            in_refs = guards.is_references((p.text or "").strip())
            found = found or in_refs
            continue
        if in_refs:
            text.append(p.text or "")
    if not found:
        return None
    return set(guards.all_citekeys("\n".join(text)))
