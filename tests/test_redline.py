"""Fixtures for the deterministic redline machinery (GPU-free).

The load-bearing assertion in this file is the equation one. A Word paragraph is NOT the
text in its w:r runs: an inline equation is an m:oMath element that is a SIBLING of the runs.
A naive "read w:t, diff, rewrite runs" approach silently deletes every equation, footnote,
field, and drawing in the paragraph. raconteur hits this immediately — a Results section is
full of inline statistics rendered as OMML.

    INVARIANT: raconteur never authors an equation; it only edits the prose around one.

So an equation must survive a sentence edit as ACCEPTED content: a direct child of w:p, never
stranded inside a w:ins or w:del.
"""
from __future__ import annotations

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from raconteur import guards, redline

_MATH = "http://schemas.openxmlformats.org/officeDocument/2006/math"


# ── helpers ───────────────────────────────────────────────────────────────────

def _para(text: str):
    return Document().add_paragraph(text)


def _omath(text: str):
    """A minimal <m:oMath> carrying `text` — an equation, as Word stores it: a sibling of the
    text runs, not inside one. Built with lxml because python-docx's OxmlElement only knows
    its own registered prefixes, and `m` is not one of them."""
    om = etree.SubElement(etree.Element("root"), f"{{{_MATH}}}oMath")
    t = etree.SubElement(etree.SubElement(om, f"{{{_MATH}}}r"), f"{{{_MATH}}}t")
    t.text = text
    return om


def _para_with_math(before: str, equation: str, after: str):
    """A paragraph whose prose is interrupted by an equation, as pandoc renders inline stats."""
    p = _para(before)
    p._p.append(_omath(equation))
    p._p.append(redline._text_run(after))
    return p


def _math_texts(p_el) -> list[str]:
    return ["".join(t.text or "" for t in om.iter(f"{{{_MATH}}}t"))
            for om in p_el.iter(f"{{{_MATH}}}oMath")]


def _accepted(p_el) -> str:
    """Text with all tracked changes accepted (insertions kept, deletions dropped)."""
    out = []
    for r in p_el.iter(qn("w:r")):
        if r.getparent().tag == qn("w:del"):
            continue
        for t in r.iter(qn("w:t")):
            out.append(t.text or "")
    return "".join(out)


def _rejected(p_el) -> str:
    """Text with all tracked changes rejected (deletions kept, insertions dropped)."""
    out = []
    for r in p_el.iter(qn("w:r")):
        if r.getparent().tag == qn("w:ins"):
            continue
        for t in r.iter(qn("w:t")):
            out.append(t.text or "")
        for t in r.iter(qn("w:delText")):
            out.append(t.text or "")
    return "".join(out)


def _n(p_el, tag: str) -> int:
    return len(p_el.findall(qn(tag)))


def _ids():
    return redline._Ids(100)


# ── the atom stream ───────────────────────────────────────────────────────────

def test_serialize_paragraph_emits_sentinel_for_equation():
    p = _para_with_math("Drift was ", "x_{t+1}", " under all conditions.")
    text = redline.paragraph_text(p._p)
    assert text == "Drift was ⟦m:1⟧ under all conditions."


def test_serialize_paragraph_plain_prose_has_no_sentinels():
    p = _para("Plain prose with no atoms.")
    assert redline.paragraph_text(p._p) == "Plain prose with no atoms."
    assert guards.sentinels(redline.paragraph_text(p._p)) == []


def test_multiple_atoms_number_sequentially():
    p = _para("A ")
    p._p.append(_omath("e1"))
    p._p.append(redline._text_run(" and "))
    p._p.append(_omath("e2"))
    p._p.append(redline._text_run(" end."))
    assert redline.paragraph_text(p._p) == "A ⟦m:1⟧ and ⟦m:2⟧ end."


# ── THE load-bearing test: an equation survives a sentence edit ───────────────

def test_equation_survives_sentence_edit_as_accepted_content():
    p = _para_with_math("Drift was ", "x_{t+1}", " under all conditions. A second sentence.")
    p_el = p._p

    old = redline.paragraph_text(p_el)
    # Edit ONLY the second sentence. The equation lives in the first.
    new = old.replace("A second sentence.", "A revised second sentence.")
    assert redline.tracked_replace_sentencewise(p_el, new, "raconteur", _ids())

    # (a) the equation still exists, exactly once, with its content intact
    assert _math_texts(p_el) == ["x_{t+1}"]

    # (b) it is a DIRECT CHILD of w:p — not stranded inside a tracked change
    omaths = p_el.findall(f"{{{_MATH}}}oMath")
    assert len(omaths) == 1, "equation is not a direct child of w:p"

    # (c) it appears inside NO w:ins and NO w:del anywhere in the tree
    for wrapper in ("w:ins", "w:del"):
        for el in p_el.iter(qn(wrapper)):
            assert not list(el.iter(f"{{{_MATH}}}oMath")), \
                f"equation was wrapped in {wrapper} — raconteur authored an equation"

    # (d) the redline reads correctly both ways
    assert "A revised second sentence." in _accepted(p_el)
    assert "A second sentence." in _rejected(p_el)


def test_untouched_sentence_is_byte_for_byte_identical():
    p = _para("First sentence stays. Second sentence changes.")
    p_el = p._p
    old = redline.paragraph_text(p_el)
    new = "First sentence stays. Second sentence is rewritten."
    assert redline.tracked_replace_sentencewise(p_el, new, "raconteur", _ids())

    # The untouched sentence must not appear in any tracked change at all.
    for wrapper in ("w:ins", "w:del"):
        for el in p_el.iter(qn(wrapper)):
            body = "".join(t.text or "" for t in el.iter(qn("w:t")))
            body += "".join(t.text or "" for t in el.iter(qn("w:delText")))
            assert "First sentence stays." not in body

    assert _accepted(p_el) == new
    assert _rejected(p_el) == old


def test_edited_sentence_produces_exactly_one_del_and_one_ins():
    p = _para("Alpha stays. Beta changes.")
    p_el = p._p
    redline.tracked_replace_sentencewise(p_el, "Alpha stays. Beta moved.", "raconteur", _ids())
    assert _n(p_el, "w:del") == 1
    assert _n(p_el, "w:ins") == 1


def test_citekeys_in_untouched_sentence_survive():
    p = _para("Grounded claim [@smith2020]. Ungrounded claim.")
    p_el = p._p
    new = "Grounded claim [@smith2020]. Now grounded [@jones2019]."
    redline.tracked_replace_sentencewise(p_el, new, "raconteur", _ids())
    assert guards.dropped_citekeys(redline.paragraph_text(p_el), _accepted(p_el)) == []
    assert "[@smith2020]" in _accepted(p_el)


def test_prose_around_an_atom_is_redlined_individually():
    # The equation sits mid-sentence; editing that sentence must keep the atom in place.
    p = _para_with_math("The rate ", "r=0.5", " was low.")
    p_el = p._p
    old = redline.paragraph_text(p_el)
    new = "The measured rate ⟦m:1⟧ was high."
    assert redline.tracked_replace_sentencewise(p_el, new, "raconteur", _ids())
    assert _math_texts(p_el) == ["r=0.5"]
    omaths = p_el.findall(f"{{{_MATH}}}oMath")
    assert len(omaths) == 1
    assert "The measured rate " in _accepted(p_el)
    assert " was high." in _accepted(p_el)


def test_dropped_atom_is_never_lost_from_the_xml():
    # If the reviser drops a sentinel, the guard rejects upstream — but the machinery must
    # still never lose the equation. Belt and braces.
    p = _para_with_math("Rate ", "r=0.5", " was low.")
    p_el = p._p
    old = redline.paragraph_text(p_el)
    new = "Rate was low."          # sentinel dropped
    assert guards.dropped_sentinels(old, new)   # the guard sees it
    redline.tracked_replace_sentencewise(p_el, new, "raconteur", _ids())
    assert _math_texts(p_el) == ["r=0.5"], "equation was lost from the document"


def test_invented_atom_renders_as_nothing():
    p = _para("Plain prose here.")
    p_el = p._p
    new = "Plain prose ⟦m:9⟧ here."   # a sentinel the original never had
    assert guards.invented_sentinels(redline.paragraph_text(p_el), new)
    redline.tracked_replace_sentencewise(p_el, new, "raconteur", _ids())
    assert _math_texts(p_el) == []      # raconteur cannot author an equation
    assert "⟦m:9⟧" not in _accepted(p_el)


# ── no-ops ────────────────────────────────────────────────────────────────────

def test_unchanged_text_is_a_noop():
    p = _para("Nothing changes here.")
    assert redline.tracked_replace_sentencewise(p._p, "Nothing changes here.", "raconteur", _ids()) is False
    assert _n(p._p, "w:ins") == 0 and _n(p._p, "w:del") == 0


def test_empty_paragraph_is_a_noop():
    p = Document().add_paragraph("")
    assert redline.tracked_replace_sentencewise(p._p, "anything", "raconteur", _ids()) is False


# ── ids ───────────────────────────────────────────────────────────────────────

def test_ids_do_not_collide_with_existing():
    doc = Document()
    p = doc.add_paragraph("Alpha stays. Beta changes.")
    ids = redline.ids_for(doc)
    redline.tracked_replace_sentencewise(p._p, "Alpha stays. Beta moved.", "raconteur", ids)
    used = {int(el.get(qn("w:id")))
            for tag in ("w:ins", "w:del") for el in p._p.iter(qn(tag))}
    assert all(v > 0 for v in used)
    assert len(used) == 2, "ins and del must get distinct ids"


# ── comment anchoring ─────────────────────────────────────────────────────────

def _add_comment_range(p, cid: str, start_idx: int, end_idx: int):
    """Bracket runs [start_idx, end_idx) of paragraph `p` with a comment range."""
    from docx.oxml import OxmlElement
    runs = p._p.findall(qn("w:r"))
    s = OxmlElement("w:commentRangeStart"); s.set(qn("w:id"), cid)
    e = OxmlElement("w:commentRangeEnd");   e.set(qn("w:id"), cid)
    runs[start_idx].addprevious(s)
    runs[end_idx - 1].addnext(e)


def test_comment_spans_and_anchored_sentences():
    # Build a paragraph whose sentences are separate runs, then anchor a comment to run 1.
    p = Document().add_paragraph("")
    p._p.append(redline._text_run("One. "))
    p._p.append(redline._text_run("Two. "))
    p._p.append(redline._text_run("Three."))
    _add_comment_range(p, "7", 1, 2)          # brackets "Two. "

    text = redline.paragraph_text(p._p)
    assert text == "One. Two. Three."

    spans = redline.comment_spans(p._p)
    assert "7" in spans
    anchored = redline.anchored_sentences(text, spans["7"])
    assert anchored == {1}, f"comment should anchor to sentence 2 only, got {anchored}"


def test_anchored_sentences_offsets_account_for_sentinels():
    # Note: findall(w:r) does not see the oMath, so run index 2 is "Next sentence.".
    # This is the good discriminator: a sentinel-blind comment_spans would start the span at
    # char 15 instead of 20, overlap sentence 0, and return {0, 1}.
    p = Document().add_paragraph("")
    p._p.append(redline._text_run("Rate "))
    p._p.append(_omath("r=0.5"))
    p._p.append(redline._text_run(" was low. "))
    p._p.append(redline._text_run("Next sentence."))
    _add_comment_range(p, "3", 2, 3)          # brackets "Next sentence."

    text = redline.paragraph_text(p._p)
    assert text == "Rate ⟦m:1⟧ was low. Next sentence."
    spans = redline.comment_spans(p._p)
    assert spans["3"][0] == len("Rate ⟦m:1⟧ was low. ")
    anchored = redline.anchored_sentences(text, spans["3"])
    assert anchored == {1}, f"should anchor to the second sentence only, got {anchored}"


def test_comment_spans_width_correct_past_the_tenth_atom():
    """Regression: hardcoding the sentinel width as len("⟦m:0⟧") drifts one char per atom
    from the tenth onward, mis-anchoring every comment after it. A Results paragraph full of
    inline statistics hits this."""
    p = Document().add_paragraph("")
    for i in range(12):                       # atoms 1..12; #10-12 are 6 chars wide
        p._p.append(_omath(f"e{i}"))
    p._p.append(redline._text_run("Tail sentence."))
    _add_comment_range(p, "9", 0, 1)          # the only w:r is the tail run

    text = redline.paragraph_text(p._p)
    expected_prefix = "".join(f"⟦m:{i}⟧" for i in range(1, 13))
    assert text == expected_prefix + "Tail sentence."
    spans = redline.comment_spans(p._p)
    assert spans["9"][0] == len(expected_prefix), (
        "comment offset drifted: sentinel width must match the emitted key"
    )


def test_comment_anchors_survive_a_redline():
    p = Document().add_paragraph("")
    p._p.append(redline._text_run("One. "))
    p._p.append(redline._text_run("Two."))
    _add_comment_range(p, "5", 1, 2)
    redline.tracked_replace_sentencewise(p._p, "One. Two revised.", "raconteur", _ids())
    assert _n(p._p, "w:commentRangeStart") == 1
    assert _n(p._p, "w:commentRangeEnd") == 1


# ── document structure ────────────────────────────────────────────────────────

@pytest.mark.parametrize("style,heading", [
    ("Heading 1", True), ("Heading 2", True), ("Title", True),
    ("Normal", False), ("Body Text", False), ("", False),
])
def test_is_heading_style(style, heading):
    assert redline.is_heading_style(style) is heading


def _structured_doc():
    doc = Document()
    doc.add_paragraph("Trust in AI", style="Title")
    doc.add_paragraph("We study trust and report an effect.")     # abstract (body prose)
    doc.add_paragraph("Background", style="Heading 1")
    doc.add_paragraph("Prior work is thin [@smith2020].")
    doc.add_paragraph("Methods", style="Heading 1")
    doc.add_paragraph("We fit a logistic model.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Smith, J. (2020). Trust.")
    return doc


def test_body_paragraphs_skips_headings_title_and_references():
    recs = redline.body_paragraphs(_structured_doc())
    texts = [r["para"].text for r in recs]
    assert texts == [
        "We study trust and report an effect.",
        "Prior work is thin [@smith2020].",
        "We fit a logistic model.",
    ]
    assert "Smith, J. (2020). Trust." not in texts, "References must never be redlined"
    assert "Trust in AI" not in texts, "the title must never be redlined"


def test_body_paragraphs_tags_section_and_kind():
    recs = redline.body_paragraphs(_structured_doc())
    # The Title is skipped but must NOT open a section — otherwise the abstract's enclosing
    # section becomes the paper's own title.
    assert [r["heading"] for r in recs] == ["", "Background", "Methods"]
    assert [r["kind"] for r in recs] == ["other", "litrev", "methods"]


def test_abstract_is_body_prose_and_is_redlined():
    # Decision: the abstract IS redlined (References are rebuilt, never redlined).
    recs = redline.body_paragraphs(_structured_doc())
    assert recs[0]["para"].text.startswith("We study trust")


def test_accepted_body_text_drops_deletions_keeps_insertions():
    doc = Document()
    p = doc.add_paragraph("Alpha stays. Beta changes.")
    redline.tracked_replace_sentencewise(p._p, "Alpha stays. Beta moved.", "raconteur", _ids())
    text = redline.accepted_body_text(doc)
    assert "Beta moved." in text
    assert "Beta changes." not in text


# ── end-to-end: a real .docx on disk, with a real Word comment ────────────────

def _annotated_paper(path):
    """A realistic annotated draft: title, abstract, sections, an inline equation in Results,
    a References list, and a genuine Word comment anchored to the last sentence."""
    doc = Document()
    doc.add_paragraph("Trust in AI", style="Title")
    doc.add_paragraph("We study trust and report a large effect.")
    doc.add_paragraph("Background", style="Heading 1")
    doc.add_paragraph("Prior work is thin [@smith2020]. It ignores reliance.")
    doc.add_paragraph("Results", style="Heading 1")
    r = doc.add_paragraph("Accuracy rose to ")
    r._p.append(_omath("0.94"))
    r._p.append(redline._text_run(" overall. A second claim needs work."))
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Smith, J. (2020). Trust.")
    doc.add_comment(runs=[r.runs[-1]], text="tighten this claim", author="DCR")
    doc.save(str(path))
    return path


def test_roundtrip_equation_and_comment_survive_save_and_reopen(tmp_path):
    """The whole point of the module, exercised against a file Word would actually open."""
    path = _annotated_paper(tmp_path / "draft.docx")

    anchors = redline.comment_anchors(path)
    assert len(anchors) == 1
    a = anchors[0]
    assert a["kind"] == "results"
    assert a["text"] == "Accuracy rose to ⟦m:1⟧ overall. A second claim needs work."
    assert redline.comments_by_id(path)[a["ids"][0]]["text"] == "tighten this claim"

    doc = Document(str(path))
    target = next(r for r in redline.body_paragraphs(doc) if r["kind"] == "results")
    p_el = target["para"]._p
    old = redline.paragraph_text(p_el)
    new = old.replace("A second claim needs work.", "A second claim is tightened.")
    assert redline.tracked_replace_sentencewise(p_el, new, "raconteur", redline.ids_for(doc))
    doc.save(str(path))

    # Reopen from disk — this is what Word sees.
    doc2 = Document(str(path))
    p2 = next(r for r in redline.body_paragraphs(doc2) if r["kind"] == "results")["para"]._p

    assert _math_texts(p2) == ["0.94"]
    assert len(p2.findall(f"{{{_MATH}}}oMath")) == 1, "equation must be a direct child of w:p"
    for wrapper in ("w:ins", "w:del"):
        for el in p2.iter(qn(wrapper)):
            assert not list(el.iter(f"{{{_MATH}}}oMath")), f"equation wrapped in {wrapper}"
    assert _n(p2, "w:commentRangeStart") == 1, "comment anchor lost across the redline"
    assert redline.paragraph_text(p2) == \
        "Accuracy rose to ⟦m:1⟧ overall. A second claim is tightened."


def test_roundtrip_references_and_title_are_never_redlined(tmp_path):
    path = _annotated_paper(tmp_path / "draft.docx")
    recs = redline.body_paragraphs(Document(str(path)))
    texts = [r["para"].text for r in recs]
    assert not any("Smith, J. (2020)" in t for t in texts)
    assert not any(t == "Trust in AI" for t in texts)


# ── flatten: the clean-rewrite path must not silently drop equations ──────────

def test_flatten_paragraph_renders_atoms_as_their_text():
    p = _para_with_math("Accuracy rose to ", "0.94", " overall.")
    assert redline.flatten_paragraph(p._p) == "Accuracy rose to 0.94 overall."


def test_python_docx_text_property_loses_the_equation():
    """The bug this guards against: paragraph.text walks w:t only, and an equation's
    characters live in m:t. --resynth read the document this way and lost every number."""
    p = _para_with_math("Accuracy rose to ", "0.94", " overall.")
    assert "0.94" not in p.text          # the naive read
    assert "0.94" in redline.flatten_paragraph(p._p)


def test_revise_read_text_preserves_equations(tmp_path):
    from raconteur import revise
    path = tmp_path / "eq.docx"
    doc = Document()
    q = doc.add_paragraph("Accuracy rose to ")
    q._p.append(_omath("0.94"))
    q._p.append(redline._text_run(" overall."))
    doc.save(str(path))
    assert "0.94" in revise.read_text(path)
