"""In-place, comment-preserving revision with tracked changes.

This is what `paper` does by default when it finds an annotated .docx. The alternative
(`paper --resynth`) re-synthesises the whole manuscript from markdown and renders a fresh
.docx — which discards the reviewer's Word comments and gives them no redline to read the
tool's edits against. This module instead edits a COPY of the annotated .docx in place: it
answers each comment by rewriting only the paragraph(s) that comment is anchored to, records
every rewrite as a Word tracked change attributed to `raconteur`, and leaves the comments
anchored and every un-flagged paragraph byte-for-byte untouched.

Two failures fall out of the clean-rewrite default, and this module exists to kill both:

  COLLATERAL DRIFT — a comment on the Discussion causes the Methods section to be rewritten.
    The reviewer sees changes they never asked for, in sections they approved. A tracked
    change that altered an untouched region, under a reply claiming success, is worse than
    no edit at all.

  NO REDLINE — the reviewer annotated a specific sentence; a clean rewrite throws the
    sentence away and writes a new paragraph. There is no way to see what changed, accept
    some edits and reject others, or trust that the parts they liked survived.

This file is the deterministic machinery only: XML surgery, GPU-free and unit-testable. The
LLM call that turns a comment into revised paragraph text lives in `revise`.

OOXML notes:
  * A comment is anchored by ``<w:commentRangeStart w:id=N/>`` … ``<w:commentRangeEnd
    w:id=N/>`` markers bracketing a run range, plus a ``<w:commentReference w:id=N/>`` run;
    the text lives in comments.xml. python-docx preserves all of these across an open/save,
    so we only manipulate the body XML.
  * A tracked deletion wraps the old run(s) in ``<w:del>`` and turns ``<w:t>`` into
    ``<w:delText>``; a tracked insertion wraps new run(s) in ``<w:ins>``. Both carry an
    author and date, and Word renders them as an accept/rejectable redline.

A paragraph is modelled as an ordered stream of TEXT and OPAQUE atoms (equations, footnote
references, drawings), NOT as the text inside its ``w:r/w:t`` runs. That older model is blind
to everything else in the paragraph: an equation is a SIBLING of the text runs, so a differ
built on ``w:t`` alone sees prose with holes where every number had been, no sentence can
match, and each rewrite collapses to a whole-paragraph replacement — with the equations left
stranded at the paragraph tail, severed from the claims they verified. raconteur hits this
immediately: a Results section is full of inline statistics rendered as OMML.

Atoms serialize to sentinels (``⟦m:1⟧``) for the differ and for the LLM, and expand back to
their original elements on write.

    INVARIANT: raconteur never authors an equation; it only edits the prose around one.

An atom is re-laid as ACCEPTED content between the redlined prose, never inside a
``w:ins``/``w:del``. ``guards.dropped_sentinels`` and ``guards.invented_sentinels`` fail the
edit closed if the model breaks the invariant.

Known limitations (documented, not bugs):
  * Multiple comments on one paragraph are coarsened to bracket the whole revised paragraph —
    every comment stays valid and anchored, but loses sub-paragraph precision. ``comment_spans``
    recovers that precision on the way IN, which is what tells the minimal-edit guard which
    sentences a comment actually bears on.
  * Assumes the annotated draft has no still-open tracked changes from a prior cycle (true for
    a freshly rendered _ra draft the reviewer annotated).
"""

from __future__ import annotations

import copy
import datetime
import difflib
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from . import guards

# OOXML math. An equation is a sibling of the text runs, NOT inside one — see module docstring.
_MATH = "http://schemas.openxmlformats.org/officeDocument/2006/math"

# python-docx's nsmap does not register the reserved ``xml`` prefix, so qn("xml:space")
# would KeyError — use the literal namespaced attribute name.
_XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"

AUTHOR = "raconteur"


def _now() -> str:
    return datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")


# ── id allocation ─────────────────────────────────────────────────────────────

class _Ids:
    """Hand out w:id values that don't collide with existing comment/change ids."""

    def __init__(self, start: int):
        self._n = start

    def next(self) -> int:
        self._n += 1
        return self._n


def _max_existing_id(doc) -> int:
    """Highest w:id already used by comments or tracked changes, body + comments part."""
    ids = [0]

    def _scan(root):
        for tag in ("w:comment", "w:commentRangeStart", "w:commentRangeEnd",
                    "w:commentReference", "w:ins", "w:del"):
            for el in root.iter(qn(tag)):
                v = el.get(qn("w:id"))
                if v and v.lstrip("-").isdigit():
                    ids.append(int(v))

    _scan(doc.element.body)
    for rel in doc.part.rels.values():
        if rel.reltype.lower().endswith("/comments"):
            _scan(rel.target_part._element)
            break
    return max(ids)


def ids_for(doc) -> _Ids:
    return _Ids(_max_existing_id(doc))


# ── element builders ──────────────────────────────────────────────────────────

def _rpr_clone(run_el):
    rpr = run_el.find(qn("w:rPr"))
    return copy.deepcopy(rpr) if rpr is not None else None


def _text_run(text: str, rpr=None):
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    t = OxmlElement("w:t")
    t.set(_XML_SPACE, "preserve")
    t.text = text
    r.append(t)
    return r


def _ins(text: str, author: str, wid: int, rpr=None):
    el = OxmlElement("w:ins")
    el.set(qn("w:id"), str(wid))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), _now())
    el.append(_text_run(text, rpr))
    return el


def _del(text: str, author: str, wid: int, rpr=None):
    el = OxmlElement("w:del")
    el.set(qn("w:id"), str(wid))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), _now())
    r = OxmlElement("w:r")
    if rpr is not None:
        r.append(copy.deepcopy(rpr))
    dt = OxmlElement("w:delText")
    dt.set(_XML_SPACE, "preserve")
    dt.text = text
    r.append(dt)
    el.append(r)
    return el


# ── the paragraph as an atom stream ───────────────────────────────────────────

_SENTINEL_SPLIT = re.compile(r"(⟦[a-z]+:\d+⟧)")

_OPAQUE_RUN_CHILDREN = ("w:footnoteReference", "w:endnoteReference",
                        "w:drawing", "w:pict", "w:object")


def _is_text_run(r) -> bool:
    """A run carrying visible text (not a comment-reference marker run)."""
    return r.find(qn("w:t")) is not None and r.find(qn("w:commentReference")) is None


def _is_ref_run(r) -> bool:
    return r.tag == qn("w:r") and r.find(qn("w:commentReference")) is not None


def _is_opaque(el) -> bool:
    """Content we preserve but never author."""
    if el.tag in (f"{{{_MATH}}}oMath", f"{{{_MATH}}}oMathPara"):
        return True
    if el.tag == qn("w:hyperlink"):
        return True
    if el.tag == qn("w:r"):
        return any(el.find(qn(t)) is not None for t in _OPAQUE_RUN_CHILDREN)
    return False


def _sentinel_kind(el) -> str:
    if el.tag.startswith(f"{{{_MATH}}}"):
        return "m"
    if el.tag == qn("w:hyperlink"):
        return "h"
    return "x"


def serialize_paragraph(p_el) -> tuple[str, dict[str, object], list]:
    """Render a paragraph as (text_with_sentinels, sentinel -> element, consumed children).

    Comment plumbing and ``w:pPr`` are left alone — they are re-attached around the rebuilt
    body. Prior tracked deletions are dropped (the paragraph reads as it currently stands);
    prior insertions read as accepted text.
    """
    parts: list[str] = []
    smap: dict[str, object] = {}
    consumed: list = []
    n = 0
    for child in list(p_el):
        tag = child.tag
        if tag == qn("w:pPr") or tag in (qn("w:commentRangeStart"), qn("w:commentRangeEnd")):
            continue
        if _is_ref_run(child) or tag == qn("w:del"):
            continue
        if _is_opaque(child):
            n += 1
            key = f"⟦{_sentinel_kind(child)}:{n}⟧"
            smap[key] = child
            parts.append(key)
            consumed.append(child)
        elif tag == qn("w:ins"):
            parts.append("".join(t.text or "" for t in child.iter(qn("w:t"))))
            consumed.append(child)
        elif tag == qn("w:r") and child.findall(qn("w:t")):
            parts.append("".join(t.text or "" for t in child.findall(qn("w:t"))))
            consumed.append(child)
    return "".join(parts), smap, consumed


def paragraph_text(p_el) -> str:
    """The paragraph as the reviser and the differ see it: prose with sentinels for atoms."""
    return serialize_paragraph(p_el)[0]


def atom_text(el) -> str:
    """The visible text inside an opaque atom.

    An equation's characters live in ``m:t``, not ``w:t``, so anything reading only ``w:t``
    sees a paragraph with holes where every number was. Used to flatten a paragraph for
    callers that want plain prose rather than sentinels.
    """
    parts = [t.text or "" for t in el.iter(f"{{{_MATH}}}t")]
    parts += [t.text or "" for t in el.iter(qn("w:t"))]
    return "".join(parts)


def flatten_paragraph(p_el) -> str:
    """The paragraph as plain prose, with each atom rendered as its own text.

    Deletions are dropped and insertions read as accepted, exactly as ``serialize_paragraph``
    defines it — but atoms come back as their characters instead of ``⟦m:1⟧``, so the result
    is readable markdown rather than a sentinel stream.
    """
    text, smap, _ = serialize_paragraph(p_el)
    for key, el in smap.items():
        text = text.replace(key, atom_text(el))
    return text


def _render(text: str, smap: dict, rpr) -> list:
    """Text with sentinels -> runs, each sentinel expanded to its original element.

    An unknown sentinel (one the model invented) renders as nothing: raconteur never authors
    an equation. The ``invented_sentinels`` guard rejects the rewrite before it reaches here,
    so this is a backstop, not a path.
    """
    out: list = []
    for piece in _SENTINEL_SPLIT.split(text):
        if not piece:
            continue
        if _SENTINEL_SPLIT.fullmatch(piece):
            el = smap.get(piece)
            if el is not None:
                out.append(copy.deepcopy(el))
        else:
            out.append(_text_run(piece, rpr))
    return out


def _segments(chunk: str) -> tuple[list[str], list[str]]:
    """Split on sentinels: (text segments, sentinels). ``len(texts) == len(sents) + 1``."""
    parts = _SENTINEL_SPLIT.split(chunk)
    return parts[0::2], parts[1::2]


def _redline_chunk(old_chunk: str, new_chunk: str, smap: dict,
                   author: str, ids: _Ids, rpr) -> list:
    """Redline one changed span, never touching an atom.

    An equation is re-laid as accepted content between the redlined prose around it —
    raconteur cannot author an equation, so it must not claim to have deleted or inserted
    one. When the sentinel sequence is unchanged (the guarded case) the prose segments around
    each atom are redlined individually, so even a rewritten sentence keeps its numbers
    exactly where they were.
    """
    o_texts, o_sents = _segments(old_chunk)
    n_texts, n_sents = _segments(new_chunk)
    body: list = []

    if o_sents != n_sents:
        # The reviser moved, dropped, or invented an atom. The guard rejects this upstream;
        # here we simply never lose one: redline the prose, then re-lay every original atom.
        o_all, n_all = "".join(o_texts), "".join(n_texts)
        if o_all:
            body.append(_del(o_all, author, ids.next(), rpr))
        if n_all:
            body.append(_ins(n_all, author, ids.next(), rpr))
        body.extend(copy.deepcopy(smap[s]) for s in o_sents if s in smap)
        return body

    for k in range(len(o_sents) + 1):
        o = o_texts[k] if k < len(o_texts) else ""
        n = n_texts[k] if k < len(n_texts) else ""
        if o and o == n:
            body.append(_text_run(o, rpr))      # unchanged prose around an atom
        else:
            if o:
                body.append(_del(o, author, ids.next(), rpr))
            if n:
                body.append(_ins(n, author, ids.next(), rpr))
        if k < len(o_sents):
            el = smap.get(o_sents[k])
            if el is not None:
                body.append(copy.deepcopy(el))  # the atom itself: accepted, in place
    return body


def _relay(p_el, body: list, consumed: list) -> None:
    """Detach what we consumed and re-lay [starts] <body> [ends] [reference runs]."""
    starts = p_el.findall(qn("w:commentRangeStart"))
    ends = p_el.findall(qn("w:commentRangeEnd"))
    ref_runs = [r for r in p_el.findall(qn("w:r")) if _is_ref_run(r)]
    for el in consumed + starts + ends + ref_runs:
        p_el.remove(el)
    ppr = p_el.find(qn("w:pPr"))
    insert_at = list(p_el).index(ppr) + 1 if ppr is not None else 0
    for offset, el in enumerate(list(starts) + body + list(ends) + list(ref_runs)):
        p_el.insert(insert_at + offset, el)


# ── tracked edits ─────────────────────────────────────────────────────────────

def tracked_replace(p_el, new_text: str, author: str = AUTHOR, ids: _Ids | None = None) -> bool:
    """Replace a paragraph's text with one tracked deletion of the old + insertion of the new,
    preserving comment anchors and every opaque atom.

    Coarse: the whole paragraph is redlined. ``tracked_replace_sentencewise`` is what the
    redline path uses; this remains for callers that genuinely mean to replace wholesale.

    Returns False (a no-op) when there is no text to replace or the text is unchanged.
    """
    ids = ids or _Ids(1000)
    old_text, smap, consumed = serialize_paragraph(p_el)
    if not consumed or new_text.strip() == old_text.strip():
        return False
    rpr = next((_rpr_clone(r) for r in consumed
                if r.tag == qn("w:r") and _is_text_run(r)), None)
    _relay(p_el, _redline_chunk(old_text, new_text, smap, author, ids, rpr), consumed)
    return True


def tracked_replace_sentencewise(p_el, new_text: str, author: str = AUTHOR,
                                 ids: _Ids | None = None) -> bool:
    """Replace a paragraph's text with SENTENCE-level tracked changes.

    Diffs old against new at sentence granularity and redlines only the sentences that
    actually changed; every unchanged sentence is re-laid as a plain (accepted) run,
    byte-for-byte, so its [@citekey] tags, its grounding, and its equations survive the
    revision untouched. Opaque atoms are never deleted or inserted — see ``_redline_chunk``.

    Returns False (a no-op) when there is no text to replace or the text is unchanged.
    """
    ids = ids or _Ids(1000)
    old_text, smap, consumed = serialize_paragraph(p_el)
    if not consumed or new_text.strip() == old_text.strip():
        return False
    rpr = next((_rpr_clone(r) for r in consumed
                if r.tag == qn("w:r") and _is_text_run(r)), None)

    old_units = guards.sentence_units(old_text)
    new_units = guards.sentence_units(new_text)
    sm = difflib.SequenceMatcher(a=old_units, b=new_units, autojunk=False)
    body: list = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for u in old_units[i1:i2]:
                body.extend(_render(u, smap, rpr))  # unchanged — accepted, atoms in place
        else:
            body.extend(_redline_chunk("".join(old_units[i1:i2]),
                                       "".join(new_units[j1:j2]),
                                       smap, author, ids, rpr))
    _relay(p_el, body, consumed)
    return True


def tracked_insert_after(p_el, text: str, author: str = AUTHOR, ids: _Ids | None = None):
    """Insert a brand-new paragraph (wholly a tracked insertion) after ``p_el``, cloning its
    paragraph properties. For structural comments that ask to split a paragraph or add
    material."""
    ids = ids or _Ids(1000)
    new_p = OxmlElement("w:p")
    ppr = p_el.find(qn("w:pPr"))
    if ppr is not None:
        new_p.append(copy.deepcopy(ppr))
    rpr = next((_rpr_clone(r) for r in p_el.findall(qn("w:r")) if _is_text_run(r)), None)
    new_p.append(_ins(text, author, ids.next(), rpr))
    p_el.addnext(new_p)
    return new_p


# ── comment reading / anchoring ───────────────────────────────────────────────

def comments_by_id(path: Path) -> dict[str, dict]:
    """Map comment id -> {author, text} from the comments part."""
    doc = Document(str(path))
    out: dict[str, dict] = {}
    for rel in doc.part.rels.values():
        if not rel.reltype.lower().endswith("/comments"):
            continue
        for c in rel.target_part._element.findall(".//" + qn("w:comment")):
            cid = c.get(qn("w:id"))
            texts = [t.text for t in c.findall(".//" + qn("w:t")) if t.text]
            out[cid] = {"author": c.get(qn("w:author"), "reviewer"),
                        "text": " ".join(texts)}
        break
    return out


def comment_spans(p_el) -> dict[str, tuple[int, int]]:
    """Character offsets ``[start, end)`` of each comment's anchored range, measured over the
    same serialized text ``paragraph_text`` returns.

    The reviewer highlights the phrase their comment is about, so this recovers WHICH
    SENTENCES a comment actually bears on. This is what lets the minimal-edit guard know that
    a comment on sentence 2 does not license rewriting sentences 1 and 3-7.
    """
    offset = 0
    n = 0
    opens: dict[str, int] = {}
    spans: dict[str, tuple[int, int]] = {}
    for child in list(p_el):
        tag = child.tag
        if tag == qn("w:commentRangeStart"):
            opens[child.get(qn("w:id"))] = offset
        elif tag == qn("w:commentRangeEnd"):
            cid = child.get(qn("w:id"))
            if cid in opens:
                spans[cid] = (opens.pop(cid), offset)
        elif tag == qn("w:pPr") or _is_ref_run(child) or tag == qn("w:del"):
            continue
        elif _is_opaque(child):
            # Must mirror serialize_paragraph's numbering exactly. Hardcoding the width as
            # len("⟦m:0⟧") drifts one char per atom from the tenth atom onward, which
            # silently mis-anchors every comment after it in a paragraph with many atoms —
            # e.g. a Results paragraph full of inline statistics.
            n += 1
            offset += len(f"⟦{_sentinel_kind(child)}:{n}⟧")
        elif tag == qn("w:ins"):
            offset += sum(len(t.text or "") for t in child.iter(qn("w:t")))
        elif tag == qn("w:r"):
            offset += sum(len(t.text or "") for t in child.findall(qn("w:t")))
    for cid, start in opens.items():  # range never closed in this paragraph
        spans[cid] = (start, offset)
    return spans


def anchored_sentences(text: str, span: tuple[int, int]) -> set[int]:
    """Indices of the sentences a comment's character range overlaps."""
    out: set[int] = set()
    pos = 0
    for i, unit in enumerate(guards.sentence_units(text)):
        if pos < span[1] and span[0] < pos + len(unit):
            out.add(i)
        pos += len(unit)
    return out


def is_heading_style(style_name: str) -> bool:
    """True for Word heading/title styles (so we never rewrite a heading as prose)."""
    s = (style_name or "").lower()
    return s.startswith("heading") or s == "title"


def is_title_style(style_name: str) -> bool:
    """The document title is not a section heading. It must be skipped like a heading, but it
    must not become the enclosing section of the abstract that follows it."""
    return (style_name or "").lower() == "title"


# ── raconteur document structure ──────────────────────────────────────────────
# rabbitHole's narrative is a flat run of body paragraphs. raconteur's .docx has a title, an
# abstract, ## section headings, ### subsection headings, and a References list. The redline
# must touch body prose only, and must know which section a paragraph belongs to so the
# reviser gets the right context bundle (a Methods sentence needs the methods writeup; a
# Background sentence needs the bib).


def _style_name(p) -> str:
    try:
        return p.style.name or ""
    except Exception:
        return ""


def body_paragraphs(doc) -> list[dict]:
    """Body prose paragraphs, each tagged with its enclosing section heading.

    Skips headings and the title (never redline a heading), everything inside a References
    section (a bibliography entry is a generated artifact, not prose the reviewer redlines),
    and empty paragraphs.

    The abstract IS included: it is body prose, and a comment on it must produce a tracked
    change like any other. Returns ``[{index, para, heading, kind}]`` in document order.
    """
    out: list[dict] = []
    heading = ""
    in_references = False
    for i, p in enumerate(doc.paragraphs):
        style = _style_name(p)
        if is_heading_style(style):
            # The title is skipped but does not open a section: the abstract that follows it
            # belongs to no section, not to a section named after the paper.
            if not is_title_style(style):
                heading = (p.text or "").strip()
                in_references = guards.is_references(heading)
            continue
        if in_references or not (p.text or "").strip():
            continue
        out.append({
            "index": i,
            "para": p,
            "heading": heading,
            "kind": guards.section_kind(heading),
        })
    return out


def comment_anchors(path: Path) -> list[dict]:
    """Body paragraphs carrying a comment anchor, with comment ids and current text.

    Returns ``[{index, para, heading, kind, ids, text, anchored}]`` in document order.
    ``text`` is the serialized paragraph (atoms as sentinels) — the exact string the reviser
    is asked to revise. ``anchored`` is the sorted set of sentence indices the paragraph's
    comments actually bear on; that set is what the minimal-edit guard enforces against.

    Paragraphs with no comment are omitted, as are headings and References — a comment on a
    heading usually means "add a section" or "find more sources", which is a routing decision,
    not a paragraph edit. The caller handles those; see ``heading_comments``.
    """
    doc = Document(str(path))
    out = []
    for rec in body_paragraphs(doc):
        p_el = rec["para"]._p
        ids = [s.get(qn("w:id")) for s in p_el.findall(qn("w:commentRangeStart"))]
        if not ids:
            continue
        text = paragraph_text(p_el)
        anchored: set[int] = set()
        for span in comment_spans(p_el).values():
            anchored |= anchored_sentences(text, span)
        out.append({**rec, "ids": ids, "text": text, "anchored": sorted(anchored)})
    return out


def heading_comments(path: Path) -> list[dict]:
    """Comments anchored to a heading. These cannot be answered by a paragraph edit — they
    ask for a section to be added, split, or resourced. The caller routes them."""
    doc = Document(str(path))
    out = []
    for i, p in enumerate(doc.paragraphs):
        if not is_heading_style(_style_name(p)):
            continue
        ids = [s.get(qn("w:id")) for s in p._p.findall(qn("w:commentRangeStart"))]
        if ids:
            out.append({"index": i, "heading": (p.text or "").strip(), "ids": ids})
    return out


def _accepted_para_text(p_el) -> str:
    """One paragraph as it reads with every tracked change accepted.

    Insertions kept, deletions dropped — ``w:t`` lives in normal and ``<w:ins>`` runs, while
    deleted text sits in ``<w:delText>`` and is therefore skipped.
    """
    out: list[str] = []
    for r in p_el.iter(qn("w:r")):
        parent = r.getparent()
        if parent is not None and parent.tag == qn("w:del"):
            continue
        for t in r.iter(qn("w:t")):
            out.append(t.text or "")
    return "".join(out)


def accepted_body_text(doc) -> str:
    """The manuscript as it reads with every tracked change accepted."""
    parts = [t for p in doc.paragraphs if (t := _accepted_para_text(p._p)).strip()]
    return "\n\n".join(parts)


def accepted_markdown(doc) -> str:
    """The post-edit manuscript rendered back to markdown, headings and all.

    ``guards`` reason over markdown (``## `` opens a section, which is how the citation floor
    gets gated on section kind and how References are excluded). Recovering that structure
    from the .docx is what lets the redline path emit the same metrics line as the draft path.
    """
    parts: list[str] = []
    for p in doc.paragraphs:
        style = _style_name(p)
        text = (p.text or "").strip()
        if is_title_style(style):
            if text:
                parts.append(f"# {text}")
            continue
        if is_heading_style(style):
            if text:
                parts.append(f"## {text}")
            continue
        accepted = _accepted_para_text(p._p).strip()
        if accepted:
            parts.append(accepted)
    return "\n\n".join(parts)
