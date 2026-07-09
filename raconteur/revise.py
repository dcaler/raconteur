from __future__ import annotations
from pathlib import Path

try:
    from docx import Document
    from docx.oxml.ns import qn
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False


def _require() -> None:
    if not _DOCX_OK:
        raise RuntimeError(
            "python-docx is required to read .docx files — pip install python-docx"
        )


def read_text(path: Path) -> str:
    """Read body text from a docx, accepting track changes (insertions visible, deletions gone).

    Reads atoms too. ``python-docx``'s ``paragraph.text`` walks ``w:t`` runs only, but an
    inline equation is an ``m:oMath`` SIBLING of those runs and its characters live in
    ``m:t`` — so the naive read returns prose with a hole where every number was, and the
    equation is silently lost from whatever is regenerated from it. ``flatten_paragraph``
    renders each atom as its own text.
    """
    _require()
    from . import redline
    doc = Document(str(path))
    parts = [t for p in doc.paragraphs if (t := redline.flatten_paragraph(p._p)).strip()]
    return "\n\n".join(parts)


def read_comments(path: Path) -> list[dict]:
    """Extract all reviewer comments from a docx."""
    _require()
    doc = Document(str(path))
    comments = []
    try:
        for rel in doc.part.rels.values():
            if "comments" not in rel.reltype.lower():
                continue
            for c in rel.target_part._element.findall(".//" + qn("w:comment")):
                author = c.get(qn("w:author"), "reviewer")
                texts = [t.text for t in c.findall(".//" + qn("w:t")) if t.text]
                if texts:
                    comments.append({"author": author, "text": " ".join(texts)})
            break
    except Exception:
        pass
    return comments


def read_track_changes(path: Path) -> dict:
    """Extract inserted and deleted text from track changes."""
    _require()
    doc = Document(str(path))
    body = doc.element.body
    insertions = []
    deletions = []
    for ins in body.iter(qn("w:ins")):
        text = "".join(t.text or "" for t in ins.iter(qn("w:t")))
        if text.strip():
            insertions.append(text)
    for dele in body.iter(qn("w:del")):
        text = "".join(t.text or "" for t in dele.iter(qn("w:delText")))
        if text.strip():
            deletions.append(text)
    return {"insertions": insertions, "deletions": deletions}


def build_revision_context(path: Path) -> str:
    """Produce a formatted summary of all annotations for the LLM."""
    comments = read_comments(path)
    changes = read_track_changes(path)

    parts = []
    if changes["deletions"]:
        lines = "\n".join(f"  - {d}" for d in changes["deletions"])
        parts.append(f"DELETED TEXT (remove these):\n{lines}")
    if changes["insertions"]:
        lines = "\n".join(f"  + {i}" for i in changes["insertions"])
        parts.append(f"INSERTED TEXT (incorporate these):\n{lines}")
    if comments:
        lines = "\n".join(f"  [{c['author']}]: {c['text']}" for c in comments)
        parts.append(f"REVIEWER COMMENTS:\n{lines}")

    return "\n\n".join(parts)
