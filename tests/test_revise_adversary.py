"""The redline must fail CLOSED.

A broken edit under a reply claiming "done" is the worst outcome — worse than a visibly
skipped comment. Malformed JSON, a dropped citekey, a dropped or invented equation, an
out-of-scope sentence, an exhausted retry budget: every one of these must leave the
reviewer's paragraph untouched and say so.

These tests drive `redline_paragraph` with a scripted brain, so they are GPU-free and
deterministic. The brain returns whatever the script says, in order.
"""
from __future__ import annotations

import pytest

from raconteur import guards
from raconteur import redline_revise as rr


class Exhausted(BaseException):
    """Deliberately NOT an Exception.

    `redline_paragraph` catches `Exception` around every brain call and fails closed. If the
    harness signalled exhaustion with an Exception, a mis-scripted test would be swallowed
    into a "skipped" outcome and look like a passing fail-closed test. This escapes that net.
    """


class ScriptedBrain:
    """Returns canned responses in order. Records every prompt for inspection."""

    def __init__(self, *responses):
        self.responses = list(responses)
        self.prompts: list[str] = []

    def coordinator(self, prompt, system="", num_ctx=0):
        self.prompts.append(prompt)
        if not self.responses:
            raise Exhausted(
                f"ScriptedBrain ran out of responses after {len(self.prompts)} call(s)")
        r = self.responses.pop(0)
        if isinstance(r, Exception):
            raise r
        return r


def _call(brain, paragraph, comments=("tighten this",), anchored={1},
          kind="litrev", known={"smith2020", "jones2019"}, rounds=2):
    return rr.redline_paragraph(
        brain, "A Paper", "Background", paragraph, list(comments),
        context_section="", bib_section="", anchored=set(anchored),
        kind=kind, known=set(known), rounds=rounds,
    )


PARA = "Grounded claim [@smith2020]. Second sentence needs work. Third stays."


# ── sentence-indexed edit machinery ──────────────────────────────────────────

def test_apply_sentence_edits_copies_untouched_byte_for_byte():
    units = guards.sentence_units(PARA)
    out = rr._apply_sentence_edits(units, {"2": "Second sentence is fixed."})
    assert out == "Grounded claim [@smith2020]. Second sentence is fixed. Third stays."
    # the untouched sentences are literally the original strings
    assert out.startswith(units[0])
    assert out.endswith(units[2])


def test_apply_sentence_edits_null_deletes():
    units = guards.sentence_units(PARA)
    out = rr._apply_sentence_edits(units, {"2": None})
    assert "Second sentence" not in out
    assert "Grounded claim [@smith2020]." in out and "Third stays." in out


def test_apply_sentence_edits_preserves_inter_sentence_spacing():
    units = guards.sentence_units("One.  Two.")
    out = rr._apply_sentence_edits(units, {"1": "Uno."})
    assert out == "Uno.  Two."


def test_number_sentences_marks_anchored():
    rendered = rr._number_sentences(guards.sentence_units(PARA), {1})
    lines = rendered.splitlines()
    assert lines[0].startswith("  1.")
    assert lines[1].startswith("▶ 2.")


@pytest.mark.parametrize("raw", [
    "not json at all",
    "```json\nbroken{\n```",
    "[1, 2, 3]",              # a list, not an object
    "",
])
def test_parse_sentence_edits_rejects_malformed(raw):
    edits, errors = rr._parse_sentence_edits(raw, 3)
    assert edits == {} and errors, "malformed output must produce an error, not a silent {}"


def test_parse_sentence_edits_rejects_out_of_range_and_bad_types():
    edits, errors = rr._parse_sentence_edits('{"9": "x", "1": 5, "2": "ok"}', 3)
    assert edits == {"2": "ok"}
    assert len(errors) == 2


def test_parse_sentence_edits_accepts_null():
    edits, errors = rr._parse_sentence_edits('{"2": null}', 3)
    assert edits == {"2": None} and not errors


def test_empty_object_is_valid_but_means_no_change():
    edits, errors = rr._parse_sentence_edits("{}", 3)
    assert edits == {} and errors == []


# ── fail-closed paths ────────────────────────────────────────────────────────

def test_malformed_json_every_round_fails_closed():
    brain = ScriptedBrain("garbage", "still garbage")
    text, outcome = _call(brain, PARA)
    assert text is None and outcome == "skipped"


def test_empty_edit_object_is_skipped_not_claimed_as_done():
    # The reviser cannot both leave the paragraph alone and have addressed the comment.
    brain = ScriptedBrain("{}")
    text, outcome = _call(brain, PARA)
    assert text is None and outcome == "skipped"


def test_dropped_citekey_fails_closed():
    # Sentence 1 carries [@smith2020]; the reviser rewrites it and drops the key.
    bad = '{"1": "Grounded claim."}'
    brain = ScriptedBrain(bad, bad)
    text, outcome = _call(brain, PARA, anchored={0})
    assert text is None and outcome == "skipped"


def test_dropped_citekey_is_reported_as_an_imperative_on_retry():
    brain = ScriptedBrain('{"1": "Grounded claim."}', '{"1": "Grounded claim [@smith2020]."}', "OK")
    text, outcome = _call(brain, PARA, anchored={0})
    assert outcome == "edited" and "[@smith2020]" in text
    # the retry prompt must have carried the dropped-citekey imperative
    assert "[@smith2020]" in brain.prompts[1]
    assert "Restore these" in brain.prompts[1]


def test_dropped_equation_fails_closed():
    para = "The rule ⟦m:1⟧ holds. Second sentence."
    bad = '{"1": "The rule holds."}'          # sentinel dropped
    brain = ScriptedBrain(bad, bad)
    text, outcome = _call(brain, para, anchored={0}, known=set())
    assert text is None and outcome == "skipped"


def test_invented_equation_fails_closed():
    para = "Plain first. Second sentence."
    bad = '{"1": "Plain ⟦m:9⟧ first."}'       # sentinel invented
    brain = ScriptedBrain(bad, bad)
    text, outcome = _call(brain, para, anchored={0}, known=set())
    assert text is None and outcome == "skipped"


def test_out_of_scope_sentence_fails_closed():
    # Comment anchors to sentence 2; the reviser also rewrites sentence 3.
    bad = '{"2": "Second fixed.", "3": "Third rewritten."}'
    brain = ScriptedBrain(bad, bad)
    text, outcome = _call(brain, PARA, anchored={1})
    assert text is None and outcome == "skipped"


def test_out_of_scope_violation_names_the_sentence_on_retry():
    brain = ScriptedBrain(
        '{"2": "Second fixed.", "3": "Third rewritten."}',
        '{"2": "Second fixed."}',
        "OK",
    )
    text, outcome = _call(brain, PARA, anchored={1})
    assert outcome == "edited"
    assert "Third stays." in text, "the out-of-scope sentence must be restored verbatim"
    assert "sentence(s) 3" in brain.prompts[1]


def test_author_year_prose_fails_closed():
    bad = '{"2": "Smith et al. (2020) found an effect."}'
    brain = ScriptedBrain(bad, bad)
    text, outcome = _call(brain, PARA, anchored={1})
    assert text is None and outcome == "skipped"


def test_unresolved_citekey_fails_closed():
    bad = '{"2": "Second sentence [@ghost99]."}'
    brain = ScriptedBrain(bad, bad)
    text, outcome = _call(brain, PARA, anchored={1})
    assert text is None and outcome == "skipped"


def test_revise_exception_fails_closed():
    brain = ScriptedBrain(RuntimeError("ollama died"))
    text, outcome = _call(brain, PARA)
    assert text is None and outcome == "skipped"


def test_audit_exception_fails_closed():
    # Guards pass, then the audit raises. We must NOT write a tracked change: the guards
    # prove the text is verifiable, not that it answers the comment.
    brain = ScriptedBrain('{"2": "Second sentence is fixed."}', RuntimeError("audit died"))
    text, outcome = _call(brain, PARA, anchored={1})
    assert text is None and outcome == "skipped"


def test_audit_rejection_then_acceptance():
    brain = ScriptedBrain(
        '{"2": "Second sentence is fixed."}',
        "1. Does not actually address the comment.",   # audit rejects
        '{"2": "Second sentence now addresses the point."}',
        "OK",
    )
    text, outcome = _call(brain, PARA, anchored={1})
    assert outcome == "edited" and "addresses the point" in text


def test_audit_rejection_exhausts_rounds_and_fails_closed():
    brain = ScriptedBrain(
        '{"2": "Attempt one."}', "1. No.",
        '{"2": "Attempt two."}', "1. Still no.",
    )
    text, outcome = _call(brain, PARA, anchored={1}, rounds=2)
    assert text is None and outcome == "skipped"


# ── the happy path ───────────────────────────────────────────────────────────

def test_clean_edit_is_accepted_and_minimal():
    brain = ScriptedBrain('{"2": "Second sentence is fixed."}', "OK")
    text, outcome = _call(brain, PARA, anchored={1})
    assert outcome == "edited"
    assert text == "Grounded claim [@smith2020]. Second sentence is fixed. Third stays."
    assert guards.touched_indices(PARA, text) == {1}


def test_methods_paragraph_needs_no_citation():
    # Section-kind gate: a Methods paragraph is grounded in the writeup, not the bibliography.
    para = "We fit a logistic model. The rate was low."
    brain = ScriptedBrain('{"2": "The learning rate was 0.01."}', "OK")
    text, outcome = _call(brain, para, anchored={1}, kind="methods")
    assert outcome == "edited" and "0.01" in text


def test_litrev_paragraph_losing_its_only_citation_fails_closed():
    para = "A claim [@smith2020]. Another sentence."
    bad = '{"1": "A claim.", "2": "Another sentence rewritten."}'
    brain = ScriptedBrain(bad, bad)
    text, outcome = _call(brain, para, anchored={0, 1}, kind="litrev")
    assert text is None and outcome == "skipped"


def test_density_guards_do_not_run_on_redline():
    """A comment asking to tighten a sentence must not license injecting citations.

    `sparse_paragraphs` would demand ceil(6/3)=2 sources here; it must not fire.
    """
    para = "A [@smith2020]. B. C. D. E. F."
    brain = ScriptedBrain('{"2": "B tightened."}', "OK")
    text, outcome = _call(brain, para, anchored={1}, kind="litrev")
    assert outcome == "edited", "a density guard fired on the redline path"


# ── routing ──────────────────────────────────────────────────────────────────

@pytest.mark.parametrize("verdict,cls", [
    ("ROUTE: section: needs a new subsection", "section"),
    ("ROUTE: sources: no such literature in the bib", "sources"),
    ("ROUTE: evidence: that ablation was never run", "evidence"),
    ("ROUTE: figure: wants a chart", "figure"),
    ("ROUTE: nonsense: unknown class", "sources"),   # unknown -> safe default
])
def test_routing_classes(verdict, cls):
    brain = ScriptedBrain('{"2": "Second sentence is fixed."}', verdict)
    text, outcome = _call(brain, PARA, anchored={1})
    assert text is None and outcome == f"route:{cls}"


def test_route_advice_is_honest_per_class():
    # Answering a figure request with "gather more sources" would be a false diagnosis.
    assert "rayleigh" in rr._ROUTE_ADVICE["evidence"]
    assert "raster" in rr._ROUTE_ADVICE["evidence"]
    assert "rabbitHole" in rr._ROUTE_ADVICE["sources"]
    assert "outline" in rr._ROUTE_ADVICE["section"]
    assert "figure" in rr._ROUTE_ADVICE["figure"] or "table" in rr._ROUTE_ADVICE["figure"]


@pytest.mark.parametrize("v,ok", [
    ("OK", True), ("ok", True), ("OK.", True), ("OK — looks good", True),
    ("1. Problem", False), ("ROUTE: section: x", False), ("", False),
])
def test_is_ok(v, ok):
    assert rr._is_ok(v) is ok


# ── naming: a redline is a minor version ─────────────────────────────────────

def test_redline_output_is_a_minor_version(tmp_path):
    user_rev = tmp_path / "260709_trust_ra_DCR.docx"
    user_rev.write_bytes(b"")
    out = rr._out_path(tmp_path, "trust", user_rev)
    assert out.name == "260709_trust_ra_DCR_ra.docx", \
        "a redline keeps the reviewer's datestamp and extends the chain"


# ── end-to-end: collateral drift is dead ─────────────────────────────────────

def _annotated_project(tmp_path, comment="tighten the second sentence"):
    """A three-section paper with an equation in Methods and one comment on Discussion."""
    from docx import Document
    from lxml import etree
    from raconteur import redline

    MATH = "http://schemas.openxmlformats.org/officeDocument/2006/math"

    def omath(t):
        om = etree.SubElement(etree.Element("root"), f"{{{MATH}}}oMath")
        x = etree.SubElement(etree.SubElement(om, f"{{{MATH}}}r"), f"{{{MATH}}}t")
        x.text = t
        return om

    pd = tmp_path / "paper"
    pd.mkdir(parents=True, exist_ok=True)
    src = pd / "260709_trust_ra_DCR.docx"

    doc = Document()
    doc.add_paragraph("Trust in AI", style="Title")
    doc.add_paragraph("Background", style="Heading 1")
    doc.add_paragraph("Prior work is thin [@smith2020]. It ignores reliance entirely.")
    doc.add_paragraph("Methods", style="Heading 1")
    m = doc.add_paragraph("We fit a logistic model with rate ")
    m._p.append(omath("0.01"))
    m._p.append(redline._text_run(" throughout."))
    doc.add_paragraph("Discussion", style="Heading 1")
    d = doc.add_paragraph("Trust is complex. This claim is vague and needs work.")
    doc.add_comment(runs=[d.runs[0]], text=comment, author="DCR")
    doc.save(str(src))
    return pd, src


def _run_redline(tmp_path, brain, comment="tighten the second sentence"):
    from raconteur.config import ProjectConfig
    pd, src = _annotated_project(tmp_path, comment)
    cfg = ProjectConfig(short_title="trust", title="Trust in AI")
    out, disp = rr.redline_revise(tmp_path, cfg, brain, pd, src,
                                  litrev="", code="", results="",
                                  bib_section="", known={"smith2020"})
    return out, disp, src


def test_e2e_only_the_commented_paragraph_changes(tmp_path):
    from docx import Document
    from docx.oxml.ns import qn
    from raconteur import redline

    brain = ScriptedBrain(
        '{"2": "This claim is now precise and supported [@smith2020]."}', "OK")
    out, disp, src = _run_redline(tmp_path, brain)

    assert out.name == "260709_trust_ra_DCR_ra.docx"
    assert disp == {"0": "edited"}

    before = {r["kind"]: redline.paragraph_text(r["para"]._p)
              for r in redline.body_paragraphs(Document(str(src)))}
    after = {r["kind"]: r for r in redline.body_paragraphs(Document(str(out)))}

    # Methods and Background are provably untouched — no tracked changes at all.
    for kind in ("methods", "litrev"):
        p_el = after[kind]["para"]._p
        assert redline.paragraph_text(p_el) == before[kind], f"{kind} drifted"
        assert len(p_el.findall(qn("w:ins"))) == 0
        assert len(p_el.findall(qn("w:del"))) == 0

    # The equation in Methods survived untouched.
    MATH = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    assert len(after["methods"]["para"]._p.findall(f"{{{MATH}}}oMath")) == 1

    # Discussion carries exactly one tracked deletion + insertion, anchor intact.
    dp = after["other"]["para"]._p
    assert len(dp.findall(qn("w:ins"))) == 1
    assert len(dp.findall(qn("w:del"))) == 1
    assert len(dp.findall(qn("w:commentRangeStart"))) == 1
    assert "Trust is complex." in redline.paragraph_text(dp)

    # The reviewer's own file is never modified.
    assert sum(len(p._p.findall(qn("w:ins"))) for p in Document(str(src)).paragraphs) == 0


def test_e2e_declined_comment_leaves_document_unchanged(tmp_path):
    from docx import Document
    from docx.oxml.ns import qn

    # Reviser drops the citekey both rounds → fail closed, no tracked change anywhere.
    bad = '{"1": "Trust is simple."}'
    brain = ScriptedBrain(bad, bad)
    out, disp, _ = _run_redline(tmp_path, brain)

    assert disp == {"0": "skipped"}
    doc = Document(str(out))
    assert sum(len(p._p.findall(qn("w:ins"))) for p in doc.paragraphs) == 0, \
        "a declined comment must leave NO tracked change"


def test_e2e_routed_comment_is_reported_not_faked(tmp_path):
    from docx import Document
    from docx.oxml.ns import qn

    # The replacement must clear the deterministic guards (Discussion expects a citation)
    # before the audit is even consulted — guards run first, by design.
    brain = ScriptedBrain('{"2": "Something [@smith2020]."}',
                          "ROUTE: evidence: that ablation was never run")
    out, disp, _ = _run_redline(tmp_path, brain, comment="run an ablation here")

    assert disp == {"0": "route:evidence"}
    doc = Document(str(out))
    assert sum(len(p._p.findall(qn("w:ins"))) for p in doc.paragraphs) == 0, \
        "a routed comment must not produce a tracked change"


def test_e2e_every_comment_gets_a_disposition(tmp_path):
    brain = ScriptedBrain('{"2": "Fixed [@smith2020]."}', "OK")
    out, disp, src = _run_redline(tmp_path, brain)
    from raconteur import redline
    cmap = redline.comments_by_id(src)
    assert set(disp) == set(cmap), "silence is not a decision"
